# smart-filtering-poc/src/smart_filtering/generator/run_jd_generation.py
import os
import sys
from pathlib import Path
import random
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[3]
SRC_DIR = PROJECT_ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from smart_filtering.config import load_config, resolve_path
from smart_filtering.generator.jd_generator import generate_jd

CONFIG = load_config()

# Define the roles for which we want to create Job Descriptions
JD_ROLES = ["Data Engineer", "Project Manager", "QA Automation Engineer"]

def create_jds_as_docx(output_dir: str, roles: list):
    """
    Generates JDs for a list of roles and saves them as formatted DOCX files.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for role in roles:
        jd_data = generate_jd(role)
        
        doc = Document()
        
        # --- Write to DOCX following the parser's expected format ---
        
        # Top-level info
        doc.add_paragraph(f"Id: {jd_data.get('id', 'N/A')}")
        doc.add_paragraph(f"Role: {jd_data.get('role', 'N/A')}")
        doc.add_paragraph(f"Min Total Years: {jd_data.get('min_total_years', 0)}")
        
        # Description
        doc.add_paragraph("### Description ###")
        doc.add_paragraph(jd_data.get('description', ''))

        # Must-Have Skills
        doc.add_paragraph("### Must Have ###")
        doc.add_paragraph(", ".join(jd_data.get('must_have', [])))

        # Nice-To-Have Skills
        doc.add_paragraph("### Nice To Have ###")
        doc.add_paragraph(", ".join(jd_data.get('nice_to_have', [])))

        # Location Policy
        doc.add_paragraph("### Location Policy ###")
        if jd_data.get('location_policy'):
            for key, value in jd_data['location_policy'].items():
                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

        # Min Skill Years
        doc.add_paragraph("### Min Skill Years ###")
        if jd_data.get('min_skill_years'):
            for skill, years in jd_data['min_skill_years'].items():
                doc.add_paragraph(f"{skill}: {years}")

        # Weights
        doc.add_paragraph("### Weights ###")
        if jd_data.get('weights'):
            for factor, weight in jd_data['weights'].items():
                doc.add_paragraph(f"{factor}: {weight}")
        
        # Save the document
        filename = f"{role.replace(' ', '_').lower()}.docx"
        file_path = os.path.join(output_dir, filename)
        doc.save(file_path)

    print(f"Successfully generated and saved {len(roles)} JDs as .docx to '{output_dir}'")

if __name__ == "__main__":
    data_cfg = CONFIG.get("data", {})
    jd_output_directory = resolve_path(data_cfg.get("jds_dir", "data/raw/jds"), project_root=PROJECT_ROOT)

    create_jds_as_docx(str(jd_output_directory), JD_ROLES)
