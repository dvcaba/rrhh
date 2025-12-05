import os
import sys
from pathlib import Path
import random
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[3]
SRC_DIR = PROJECT_ROOT / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from smart_filtering.config import load_config, resolve_path
from smart_filtering.generator.cv_generator import generate_cv, CORE_SKILLS_BY_ROLE

CONFIG = load_config()

def create_cvs_as_docx(output_dir: str, num_cvs: int = 15):
    """
    Generates a specified number of CVs and saves them as formatted DOCX files.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    available_roles = list(CORE_SKILLS_BY_ROLE.keys())
    # Build a deterministic schedule to guarantee buenos perfiles por rol: high/med/low
    role_relevance_schedule = []
    for role in available_roles:
        role_relevance_schedule.extend([(role, 2), (role, 1), (role, 0)])
    if not role_relevance_schedule:
        role_relevance_schedule = [("Generic", 1)]

    for i in range(num_cvs):
        target_role, relevance = role_relevance_schedule[i % len(role_relevance_schedule)]
        cv_data = generate_cv(target_role=target_role, relevance_hint=relevance)
        _write_cv_docx(cv_data, output_dir)

    # Add a few intentionally weak CVs (KO) to garantizar candidatos rechazados
    for role in available_roles:
        bad_cv = generate_cv(target_role=role, relevance_hint=-1)
        bad_cv["skills"] = {}  # sin must-have
        bad_cv["experience_years_total"] = round(random.uniform(0.1, 1.0), 1)
        bad_cv["experiences"] = []
        bad_cv["education"] = []
        bad_cv["languages"] = {}
        bad_cv["certs"] = []
        _write_cv_docx(bad_cv, output_dir)
    print(f"Successfully generated and saved {num_cvs + len(available_roles)} CVs as .docx to '{output_dir}'")


def _write_cv_docx(cv_data: dict, output_dir: str) -> None:
    """Helper to write a single CV dict to DOCX."""
    doc = Document()
    
    # --- Write to DOCX following the parser's expected format ---
    
    # Top-level info
    doc.add_paragraph(f"Id: {cv_data.get('id', 'N/A')}")
    doc.add_paragraph(f"Name: {cv_data.get('name', 'N/A')}")
    doc.add_paragraph(f"Title: {cv_data.get('title', 'N/A')}")
    doc.add_paragraph(f"Experience Years Total: {cv_data.get('experience_years_total', 0)}")
    doc.add_paragraph(f"Remote Preference: {cv_data.get('remote_preference', 'N/A')}")
    if cv_data.get('location'):
        doc.add_paragraph(f"Location City: {cv_data['location'].get('city', 'N/A')}")
        doc.add_paragraph(f"Location Country: {cv_data['location'].get('country', 'N/A')}")
        
    # Skills section
    doc.add_paragraph("\n### Skills ###")
    if cv_data.get('skills'):
        for skill, level in cv_data['skills'].items():
            doc.add_paragraph(f"{skill}: {level}")

    # Experience section
    doc.add_paragraph("\n### Experience ###")
    if cv_data.get('experiences'):
        for i, exp in enumerate(cv_data['experiences']):
            doc.add_paragraph(f"Role: {exp.get('role', 'N/A')}")
            doc.add_paragraph(f"Company: {exp.get('company', 'N/A')}")
            doc.add_paragraph(f"Years: {exp.get('years', 'N/A')}")
            doc.add_paragraph(f"Start Date: {exp.get('start_date', 'N/A')}")
            doc.add_paragraph(f"End Date: {exp.get('end_date', 'N/A')}")
            if exp.get('skills'):
                doc.add_paragraph(f"Skills: {', '.join(exp['skills'])}")
            if i < len(cv_data['experiences']) - 1:
                doc.add_paragraph("---") # Separator

    # Education section
    doc.add_paragraph("\n### Education ###")
    if cv_data.get('education'):
        for edu in cv_data['education']:
            doc.add_paragraph(edu)

    # Languages section
    doc.add_paragraph("\n### Languages ###")
    if cv_data.get('languages'):
        for lang, level in cv_data['languages'].items():
            doc.add_paragraph(f"{lang}: {level}")
            
    # Certifications section
    doc.add_paragraph("\n### Certifications ###")
    if cv_data.get('certs'):
        for cert in cv_data['certs']:
            doc.add_paragraph(cert)
    
    # Save the document
    file_path = os.path.join(output_dir, f"{cv_data['id']}.docx")
    doc.save(file_path)


if __name__ == "__main__":
    data_cfg = CONFIG.get("data", {})
    cv_output_directory = resolve_path(data_cfg.get("cvs_dir", "data/raw/cvs"), project_root=PROJECT_ROOT)

    create_cvs_as_docx(str(cv_output_directory), 15)
