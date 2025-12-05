# src/parser/docx_parser.py
import re
from typing import Dict, Any, List

from docx import Document

from smart_filtering.generator.cv_generator import CITIES
from smart_filtering.normalizer.skills_taxonomy import get_canonical_skill

def parse_experience(text_block: str) -> Dict[str, Any]:
    """Parses a text block representing one work experience."""
    exp = {}
    lines = text_block.strip().split('\n')
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            # Normalize key to match the dictionary keys used in the application
            key = key.strip().lower().replace(' ', '_')
            value = value.strip()
            if key == 'skills':
                exp[key] = [get_canonical_skill(s.strip()) for s in value.split(',')]
            else:
                exp[key] = value
    return exp

def parse_docx_cv(file_path: str) -> Dict[str, Any]:
    """
    Parses a DOCX CV file and reconstructs the structured dictionary.
    """
    try:
        document = Document(file_path)
        # Rebuild the text content, preserving paragraph breaks
        text = "\n".join([para.text for para in document.paragraphs if para.text])
    except Exception as e:
        print(f"Error reading or parsing DOCX file {file_path}: {e}")
        return {}

    # Initialize with default structure
    cv_data = {
        "experiences": [], 
        "skills": {},
        "education": [], 
        "languages": {},
        "certs": [],
        "location": {}
    }
    
    header_part = text.split('###')[0]
    top_level_keys = ["id", "name", "title", "experience_years_total", "remote_preference", "location_city", "location_country"]
    
    for line in header_part.split('\n'):
        if ':' in line:
            key, value = line.split(':', 1)
            key_norm = key.strip().lower().replace(' ', '_')
            value_norm = value.strip()
            
            if key_norm in top_level_keys:
                if key_norm == 'location_city':
                    cv_data["location"]["city"] = value_norm
                elif key_norm == 'location_country':
                    cv_data["location"]["country"] = value_norm
                else:
                    cv_data[key_norm] = value_norm

    sections = re.split(r'###\s*(.*?)\s*###', text)
    
    for i in range(1, len(sections), 2):
        section_title = sections[i].strip().lower()
        section_content = sections[i+1].strip()
        
        if section_title == 'skills':
            for line in section_content.split('\n'):
                if ':' in line:
                    skill, level = line.split(':', 1)
                    canonical_skill = get_canonical_skill(skill.strip())
                    cv_data["skills"][canonical_skill] = level.strip()
        
        elif section_title == 'education':
            cv_data["education"] = [edu.strip() for edu in section_content.split('\n') if edu.strip()]

        elif section_title == 'languages':
            for line in section_content.split('\n'):
                if ':' in line:
                    lang, level = line.split(':', 1)
                    cv_data["languages"][lang.strip()] = level.strip()

        elif section_title == 'certifications':
            cv_data["certs"] = [cert.strip() for cert in section_content.split('\n') if cert.strip()]
            
        elif section_title == 'experience':
            experience_blocks = section_content.split('---')
            for block in experience_blocks:
                if block.strip():
                    parsed_exp = parse_experience(block)
                    if parsed_exp:
                        cv_data["experiences"].append(parsed_exp)

    if 'experience_years_total' in cv_data and cv_data['experience_years_total']:
        try:
            cv_data['experience_years_total'] = float(cv_data['experience_years_total'])
        except (ValueError, TypeError):
            cv_data['experience_years_total'] = 0.0

    # Enrich location with coordinates when the city is known
    city = cv_data.get("location", {}).get("city")
    if city:
        city_match = next((c for c in CITIES if c["city"].lower() == city.lower()), None)
        if city_match:
            cv_data["location"].update({"lat": city_match["lat"], "lon": city_match["lon"]})
            
    return cv_data


def parse_docx_jd(file_path: str) -> Dict[str, Any]:
    """
    Parses a DOCX JD file and reconstructs the structured dictionary.
    """
    try:
        document = Document(file_path)
    except Exception as e:
        print(f"Error reading or parsing DOCX JD file {file_path}: {e}")
        return {}

    jd_data = {
        "must_have": [],
        "nice_to_have": [],
        "min_skill_years": {},
        "weights": {},
        "location_policy": {},
        "description": ""
    }
    
    current_section = None
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for para in paragraphs:
        if para.startswith('###') and para.endswith('###'):
            current_section = para.strip('# ').strip().lower()
            continue

        if current_section is None:
            if ':' in para:
                key, value = para.split(':', 1)
                key_norm = key.strip().lower().replace(' ', '_')
                value = value.strip()
                if key_norm in ['id', 'role']:
                    jd_data[key_norm] = value
                elif key_norm == 'min_total_years':
                    try:
                        jd_data[key_norm] = int(value)
                    except ValueError:
                        jd_data[key_norm] = 0
        
        elif current_section == 'description':
            jd_data['description'] += (' ' if jd_data['description'] else '') + para

        elif current_section == 'must have':
            jd_data['must_have'].extend([item.strip() for item in para.split(',')])

        elif current_section == 'nice to have':
            jd_data['nice_to_have'].extend([item.strip() for item in para.split(',')])

        elif current_section == 'location policy':
            if ':' in para:
                key, value = para.split(':', 1)
                key_norm = key.strip().lower().replace(' ', '_')
                value = value.strip()
                jd_data['location_policy'][key_norm] = value

        elif current_section == 'min skill years':
            if ':' in para:
                key, value = para.split(':', 1)
                jd_data['min_skill_years'][key.strip()] = int(value.strip())

        elif current_section == 'weights':
            if ':' in para:
                key, value = para.split(':', 1)
                jd_data['weights'][key.strip()] = float(value.strip())

    if 'max_km' in jd_data.get('location_policy', {}):
        try:
            jd_data['location_policy']['max_km'] = int(jd_data['location_policy']['max_km'])
        except (ValueError, TypeError):
            pass

    jd_data["embeddings"] = {"jd_vec": []}
            
    return jd_data
