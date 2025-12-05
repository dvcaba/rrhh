from pathlib import Path

from docx import Document

from smart_filtering.parser.docx_parser import parse_docx_cv, parse_docx_jd


def _write_cv_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Id: cv_test")
    doc.add_paragraph("Name: Test User")
    doc.add_paragraph("Title: Data Engineer")
    doc.add_paragraph("Experience Years Total: 5")
    doc.add_paragraph("Remote Preference: remote")
    doc.add_paragraph("Location City: Madrid")
    doc.add_paragraph("Location Country: ES")

    doc.add_paragraph("### Skills ###")
    doc.add_paragraph("Python: advanced")
    doc.add_paragraph("SQL: intermediate")

    doc.add_paragraph("### Experience ###")
    doc.add_paragraph("Role: Data Engineer")
    doc.add_paragraph("Company: ExampleCorp")
    doc.add_paragraph("Years: 3")
    doc.add_paragraph("Start Date: 2020-01-01")
    doc.add_paragraph("End Date: 2023-01-01")
    doc.add_paragraph("Skills: python, sql")
    doc.add_paragraph("---")

    doc.add_paragraph("### Education ###")
    doc.add_paragraph("BSc Computer Science")

    doc.add_paragraph("### Languages ###")
    doc.add_paragraph("es: native")

    doc.add_paragraph("### Certifications ###")
    doc.add_paragraph("AWS Certified Solutions Architect")

    path = tmp_path / "cv_test.docx"
    doc.save(path)
    return path


def _write_jd_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Id: jd_test")
    doc.add_paragraph("Role: Data Engineer")
    doc.add_paragraph("Min Total Years: 3")

    doc.add_paragraph("### Description ###")
    doc.add_paragraph("Data Engineer role working with Python and SQL.")

    doc.add_paragraph("### Must Have ###")
    doc.add_paragraph("python, sql")

    doc.add_paragraph("### Nice To Have ###")
    doc.add_paragraph("airflow")

    doc.add_paragraph("### Location Policy ###")
    doc.add_paragraph("Type: remote")
    doc.add_paragraph("City: Madrid")
    doc.add_paragraph("Max Km: 0")

    doc.add_paragraph("### Min Skill Years ###")
    doc.add_paragraph("python: 2")

    doc.add_paragraph("### Weights ###")
    doc.add_paragraph("skill_semantic: 0.5")
    doc.add_paragraph("title_semantic: 0.2")
    doc.add_paragraph("experience: 0.2")
    doc.add_paragraph("location: 0.1")

    path = tmp_path / "jd_test.docx"
    doc.save(path)
    return path


def test_parse_docx_cv(tmp_path: Path):
    cv_path = _write_cv_docx(tmp_path)
    parsed = parse_docx_cv(str(cv_path))

    assert parsed["id"] == "cv_test"
    assert parsed["name"] == "Test User"
    assert parsed["skills"]["python"] == "advanced"
    assert parsed["location"]["city"].lower() == "madrid"
    assert parsed["location"]["lat"] and parsed["location"]["lon"]
    assert parsed["experiences"], "Expected at least one experience parsed"


def test_parse_docx_jd(tmp_path: Path):
    jd_path = _write_jd_docx(tmp_path)
    parsed = parse_docx_jd(str(jd_path))

    assert parsed["id"] == "jd_test"
    assert "python" in parsed["must_have"]
    assert parsed["min_total_years"] == 3
    assert parsed["weights"]["skill_semantic"] == 0.5
